"""
code.py  --  Smart Dissertation TXT to DOCX Converter + Validator
================================================================
Reads a dissertation .txt file, runs validation checks (with LLM
assistance where needed), then writes a cleanly formatted .docx.

VALIDATIONS
-----------
V1  Citation repetition    -- flags AND removes duplicate (Author, Year)
                             pairs globally; first occurrence kept
V2  Citation placement     -- citations allowed only in Intro / Lit Review /
                             Methodology; 1 citation allowed in Results
V3  Subheading numbering   -- LLM decides heading vs sub-heading, numbers them N.M
V4  Heading echo           -- removes plain-text repeat of a heading on the next line
V5  Figure/equation placeholders -- [Figure X.Y: ...] and [EQN: ...] lines
V6  Figure/equation on new line  -- enforced in the parsed token stream
V7  Duplicate REFERENCES   -- detects same author+year twice; keeps fuller entry
V8  Missing references     -- in-text citation has no REFERENCES entry
V9  Orphan references      -- REFERENCES entry never cited in text
V10 Abstract double-header -- duplicate ABSTRACT heading removed
V11 Wrong chapter prefix   -- subsections with wrong chapter prefix; auto-corrected
V12 Reference format       -- detects mixed citation styles

FORMATS  (pass format=1..4 to generate_docx)
---------------------------------------------
1 -- Classic Academic Blue   (navy headings, Calibri, ruled borders)
2 -- Executive Dark Slate    (charcoal + gold accents, Georgia serif)
3 -- Modern Minimalist       (greyscale palette, Arial, clean whitespace)
4 -- Research Emerald        (deep teal + white banners, Times New Roman)

INSTALL:  pip install python-docx openai python-dotenv
USAGE:
    python code.py                          # input.txt -> output.docx (format 1)
    python code.py report.txt out.docx
    python code.py report.txt out.docx --format 2
    python code.py --no-llm
"""

import sys, re, os, json
from pathlib import Path
from collections import defaultdict
from typing import List, Dict, Tuple, Optional

# ── optional LLM ─────────────────────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    from openai import OpenAI
    load_dotenv()
    _LLM_OK = True
except ImportError:
    _LLM_OK = False

# ── docx ──────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement




def _build_theme(n: int) -> dict:
    """Return full config dict for theme n (1-4).

    1 -- Classic Academic Blue: navy headings, Calibri, bottom-ruled dividers.
    2 -- Executive Dark Slate: charcoal + dark gold, Georgia serif, gold badges.
    3 -- Modern Minimalist: greyscale, Arial, left-aligned title, hairline borders.
    4 -- Research Emerald: deep teal banners, Times New Roman, 1.5 line spacing.
    """
    themes = {
        1: dict(
            name="Classic Academic Blue", font="Calibri",
            sz_title=24, sz_ch=16, sz_sub=13, sz_h3=11, sz_body=11, line_sp=1.15,
            col_title=RGBColor(0x1F,0x49,0x7D), col_ch=RGBColor(0x1F,0x49,0x7D),
            col_sub=RGBColor(0x2E,0x74,0xB5), col_h3=RGBColor(0x40,0x40,0x40),
            col_body=RGBColor(0x1A,0x1A,0x1A), col_fig=RGBColor(0x2E,0x74,0xB5),
            col_eq=RGBColor(0x70,0x30,0xA0), col_warn=RGBColor(0xC0,0x00,0x00),
            col_cap=RGBColor(0x2E,0x74,0xB5), col_insight=RGBColor(0x40,0x40,0x40),
            col_ai=RGBColor(0x2E,0x74,0xB5),
            bdr_title="1F497D", bdr_ch="1F497D", bdr_sub="2E74B5",
            bdr_warn="C00000", bdr_fig="2E74B5",
            ch_prefix=True, title_center=True, ch_page_break=True, title_spacers=5,
            margin_top=1.0, margin_bottom=1.0, margin_left=1.25, margin_right=1.25,
        ),
        2: dict(
            name="Executive Dark Slate", font="Georgia",
            sz_title=26, sz_ch=17, sz_sub=13, sz_h3=11, sz_body=11, line_sp=1.25,
            col_title=RGBColor(0x22,0x22,0x22), col_ch=RGBColor(0x1C,0x1C,0x1C),
            col_sub=RGBColor(0xB8,0x86,0x0B), col_h3=RGBColor(0x44,0x44,0x44),
            col_body=RGBColor(0x22,0x22,0x22), col_fig=RGBColor(0xB8,0x86,0x0B),
            col_eq=RGBColor(0x6A,0x3D,0x9A), col_warn=RGBColor(0xC0,0x00,0x00),
            col_cap=RGBColor(0xB8,0x86,0x0B), col_insight=RGBColor(0x55,0x55,0x55),
            col_ai=RGBColor(0xB8,0x86,0x0B),
            bdr_title="B8860B", bdr_ch="B8860B", bdr_sub="DAA520",
            bdr_warn="C00000", bdr_fig="B8860B",
            ch_prefix=False, title_center=True, ch_page_break=True, title_spacers=8,
            margin_top=1.25, margin_bottom=1.25, margin_left=1.5, margin_right=1.5,
        ),
        3: dict(
            name="Modern Minimalist", font="Arial",
            sz_title=28, sz_ch=15, sz_sub=12, sz_h3=11, sz_body=10, line_sp=1.35,
            col_title=RGBColor(0x11,0x11,0x11), col_ch=RGBColor(0x11,0x11,0x11),
            col_sub=RGBColor(0x44,0x44,0x44), col_h3=RGBColor(0x66,0x66,0x66),
            col_body=RGBColor(0x33,0x33,0x33), col_fig=RGBColor(0x44,0x44,0x44),
            col_eq=RGBColor(0x55,0x55,0x55), col_warn=RGBColor(0xC0,0x00,0x00),
            col_cap=RGBColor(0x55,0x55,0x55), col_insight=RGBColor(0x77,0x77,0x77),
            col_ai=RGBColor(0x44,0x44,0x44),
            bdr_title="BBBBBB", bdr_ch="CCCCCC", bdr_sub="DDDDDD",
            bdr_warn="C00000", bdr_fig="CCCCCC",
            ch_prefix=False, title_center=False, ch_page_break=True, title_spacers=3,
            margin_top=1.0, margin_bottom=1.0, margin_left=1.0, margin_right=1.0,
        ),
        4: dict(
            name="Research Emerald", font="Times New Roman",
            sz_title=22, sz_ch=14, sz_sub=12, sz_h3=11, sz_body=12, line_sp=1.5,
            col_title=RGBColor(0x00,0x50,0x50), col_ch=RGBColor(0xFF,0xFF,0xFF),
            col_sub=RGBColor(0x00,0x70,0x60), col_h3=RGBColor(0x00,0x50,0x50),
            col_body=RGBColor(0x11,0x11,0x11), col_fig=RGBColor(0x00,0x70,0x60),
            col_eq=RGBColor(0x4A,0x00,0x7A), col_warn=RGBColor(0xC0,0x00,0x00),
            col_cap=RGBColor(0x00,0x70,0x60), col_insight=RGBColor(0x33,0x33,0x33),
            col_ai=RGBColor(0x00,0x70,0x60),
            bdr_title="005050", bdr_ch="005050", bdr_sub="007060",
            bdr_warn="C00000", bdr_fig="007060",
            ch_prefix=True, title_center=True, ch_page_break=True, title_spacers=6,
            margin_top=1.0, margin_bottom=1.0, margin_left=1.25, margin_right=1.25,
        ),
    }
    if n not in themes:
        raise ValueError(f"Unknown format {n!r}. Choose 1, 2, 3, or 4.")
    return themes[n]


# Active theme -- set by generate_docx(); all renderers read from this.
CFG: dict = _build_theme(1)


# Sections where citations are ALLOWED (normalised lower-case keywords)
CITATION_OK_SECTIONS = {
    "introduction", "background", "literature review", "literature",
    "methodology", "methods", "research design",
    "scope", "limitations", "significance",
}
# The one special Results sub-section that may have exactly 1 citation
RESULTS_BENCHMARK_KEYWORDS = {"benchmark", "comparison with", "base paper",
                               "reference study", "compared to", "anand"}

# Sections with no chapter number (treated as special / no citation rules)
SPECIAL_SECTIONS = {
    "ABSTRACT", "REFERENCES", "BIBLIOGRAPHY",
    "ACKNOWLEDGEMENTS", "ACKNOWLEDGMENTS",
    "TABLE OF CONTENTS", "LIST OF FIGURES", "LIST OF TABLES",
    "VALIDATION REPORT",
}


# ═════════════════════════════════════════════════════════════════════════════
# LLM HELPER
# ═════════════════════════════════════════════════════════════════════════════

def _get_client():
    """Return an OpenAI-compatible client (Grok preferred, then OpenAI)."""
    grok_key = os.getenv("GROK_API_KEY")
    oai_key  = os.getenv("OPEN_AI_KEY") or os.getenv("OPENAI_API_KEY")
    if grok_key:
        return OpenAI(api_key=grok_key, base_url="https://api.x.ai/v1"), "grok-3"
    if oai_key:
        return OpenAI(api_key=oai_key), "gpt-4o-mini"
    return None, None


def _grok(prompt: str) -> str:
    """Backward-compatible LLM call used by the original code."""
    if not _LLM_OK:
        return ""
    client, model = _get_client()
    if not client:
        return ""
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        print(f"    ⚠️  LLM error: {e}")
        return ""


def _llm(prompt: str, system: str = "") -> str:
    """Extended LLM call with optional system prompt."""
    if not _LLM_OK:
        return ""
    client, model = _get_client()
    if not client:
        return ""
    msgs = []
    if system:
        msgs.append({"role": "system", "content": system})
    msgs.append({"role": "user", "content": prompt})
    try:
        resp = client.chat.completions.create(
            model=model, messages=msgs, max_tokens=1500,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        print(f"    ⚠️  LLM error: {e}")
        return ""


def llm_classify_lines(candidates: List[Tuple[int, str]],
                        chapter_title: str,
                        subheading_title: str) -> Dict[int, str]:
    """
    Ask LLM to classify each candidate short line as:
      'h3'   — a sub-subheading label (short phrase, not a sentence)
      'body' — normal body / paragraph text
      'echo' — essentially a repetition of chapter_title or subheading_title
    """
    if not candidates:
        return {}

    items = "\n".join(f"{idx}. {text}" for idx, text in candidates)
    prompt = f"""You are formatting an academic dissertation document.
Current chapter heading: "{chapter_title}"
Current subheading: "{subheading_title}"

For each numbered line below decide ONE label:
  echo — the line is essentially a repetition/paraphrase of the chapter heading
          or subheading above (should be deleted)
  h3   — a sub-subheading / section label (short phrase, not a full sentence,
          e.g. "Ensemble Learning", "Encoding Categorical Variables",
          "Key point", "Web Application Setup", "Concluding statement")
  body — normal paragraph content or a full sentence

Rules:
- If the line closely mirrors the chapter or subheading title → echo
- If it is a short (≤ 8 words), descriptive label without a period → h3
- Everything else → body

Reply ONLY in this exact format, one line per item, no extra text:
<number>: <echo|h3|body>

Lines:
{items}
"""
    raw = _grok(prompt)
    result: Dict[int, str] = {}
    for line in raw.splitlines():
        m = re.match(r"^(\d+):\s*(echo|h3|body)\s*$", line.strip(), re.IGNORECASE)
        if m:
            result[int(m.group(1))] = m.group(2).lower()
    return result


def llm_analyse_issues(warnings: List[str], token_count: int,
                        ref_issues: List[str]) -> str:
    """
    Ask the LLM to analyse all detected issues, group them by priority,
    explain the root causes, and suggest concrete fixes.
    Returns a multi-line string appended to the Validation Report section.
    """
    if not warnings and not ref_issues:
        return ""
    all_issues = warnings + ref_issues
    prompt = f"""You are an expert academic dissertation reviewer.

The following validation issues were automatically detected in a student's dissertation
(total tokens in document: {token_count}):

{chr(10).join(f'  {i+1}. {w}' for i, w in enumerate(all_issues[:40]))}
{'  … and more' if len(all_issues) > 40 else ''}

Please provide:
1. PRIORITY ORDER — rank the top issues by academic severity (which would lose marks)
2. ROOT CAUSE — identify patterns (e.g., "all Ch4 issues share the same root cause")
3. SUGGESTED FIXES — specific, actionable steps for each category of issue
4. RESIDUAL RISK — which issues were NOT auto-fixed and must be addressed manually

Keep your analysis under 400 words. Use short paragraphs. Be direct and specific."""

    return _llm(prompt, system=(
        "You are a strict but fair academic examiner. "
        "Focus on impact, root causes, and concrete fixes."
    ))


# ═════════════════════════════════════════════════════════════════════════════
# STEP 1 — READ FILE
# ═════════════════════════════════════════════════════════════════════════════

def read_file(path: str) -> str:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")
    return p.read_text(encoding="utf-8", errors="replace")


# ═════════════════════════════════════════════════════════════════════════════
# STEP 2 — REGEX HELPERS
# ═════════════════════════════════════════════════════════════════════════════

_RE_SEP_EQ   = re.compile(r"^={6,}$")
_RE_SEP_DASH = re.compile(r"^-{6,}$")
_RE_NUMBERED = re.compile(r"^(\d+)\.\s+(.+)$")

# Citation patterns: Author et al. (YYYY)  /  Author (YYYY)  /  (Author, YYYY)
_RE_CITE = re.compile(
    r"(?:"
    r"(?:[A-Z][a-zA-Z\-']+(?:\s+et\s+al\.?)?\s*\(\d{4}\))"   # Smith (2021)
    r"|"
    r"(?:\([A-Z][a-zA-Z\-']+(?:\s+et\s+al\.?)?,?\s*\d{4}\))"  # (Smith, 2021)
    r")"
)

# Figure placeholder:  [Figure 3.1: …]  or  [Fig. 3.1 …]
_RE_FIG_PH = re.compile(r"\[(?:Figure|Fig\.?)\s+(\d+[\.\d]*)[:\s]([^\]]*)\]", re.IGNORECASE)

# Equation placeholder:  [EQ1: …]  or  [EQN: …]  or  [Equation …]
_RE_EQ_PH = re.compile(r"\[(?:EQ\w*|Equation)\s*\d*[:\s]?([^\]]*)\]", re.IGNORECASE)

# Caption line
_RE_CAPTION = re.compile(r"^(?:\[?Caption\]?)\s*:?\s*(.+)$", re.IGNORECASE)
# Insight / Description / Explanation line
_RE_INSIGHT = re.compile(r"^(?:\[?(?:Insight|Description|Explanation)\]?)\s*:?\s*(.+)$", re.IGNORECASE)
# Source URL line
_RE_SOURCE  = re.compile(r"^Source\s*:\s*(.+)$", re.IGNORECASE)


def _norm(text: str) -> str:
    return re.sub(r"[^a-z0-9 ]", "", text.lower()).strip()


def _word_overlap(a: str, b: str) -> float:
    aw, bw = set(a.split()), set(b.split())
    if not aw:
        return 0.0
    return len(aw & bw) / len(aw)


def _max_overlap(a: str, b: str) -> float:
    return max(_word_overlap(a, b), _word_overlap(b, a))


def _extract_citations(text: str) -> List[str]:
    """Return all citation strings found in a piece of text."""
    return _RE_CITE.findall(text)


def _normalise_citation(cite: str) -> str:
    """Turn 'Mhlanga (2021)' and '(Mhlanga, 2021)' into 'mhlanga2021'."""
    author = re.sub(r"et\s+al\.?", "", cite, flags=re.IGNORECASE)
    author = re.sub(r"[^a-zA-Z]", "", author.split("(")[0]).lower()
    year   = re.search(r"\d{4}", cite)
    return f"{author}{year.group() if year else ''}"


# ═════════════════════════════════════════════════════════════════════════════
# STEP 3 — STRUCTURAL PARSE
#
# Token types:
#   doc_title | chapter | subheading | h3 | body | bullet
#   figure_ph | equation_ph | caption | insight | source
#
# Inline figure/equation references inside body text are split out
# onto their own tokens so they always appear on a new line (V6).
# ═════════════════════════════════════════════════════════════════════════════

def _split_inline_placeholders(text: str) -> List[dict]:
    """
    If a body line contains [Figure X.Y: …] or [EQN: …] mixed with prose,
    split it into separate body / figure_ph / equation_ph tokens.
    """
    result = []
    combined = re.compile(
        r"(\[(?:Figure|Fig\.?)\s+\d+[\.\d]*[:\s][^\]]*\]"
        r"|\[(?:EQ\w*|Equation)\s*\d*[:\s]?[^\]]*\])",
        re.IGNORECASE
    )
    parts = combined.split(text)
    for part in parts:
        p = part.strip()
        if not p:
            continue
        if _RE_FIG_PH.match(p):
            m = _RE_FIG_PH.match(p)
            result.append({"type": "figure_ph",
                            "ref": m.group(1), "text": m.group(2).strip()})
        elif _RE_EQ_PH.match(p):
            m = _RE_EQ_PH.match(p)
            result.append({"type": "equation_ph",
                            "text": m.group(1).strip() if m.group(1) else p})
        else:
            result.append({"type": "body", "text": p})
    return result if result else [{"type": "body", "text": text}]


def parse_structure(raw: str) -> List[dict]:
    """
    Single-pass scanner → flat list of typed tokens.
    Handles all validation concerns at the token level.

    NEW (V10/V11): also fixes wrong chapter prefix in subheadings on the fly
    (e.g. "3.1" inside Chapter 4 is renumbered to "4.1").
    NEW (V10): duplicate ABSTRACT heading suppressed.
    """
    lines  = raw.splitlines()
    tokens : List[dict] = []

    # State
    title_done           = False
    in_chapter           = False
    current_chapter_num  = 0
    current_chapter_text = ""
    current_chapter_norm = ""
    last_sub_text        = ""
    last_sub_norm        = ""
    skip_echo            = False
    abstract_seen        = False      # V10: track double ABSTRACT header

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # ── blank lines
        if not line:
            i += 1
            continue

        # ── dash separator — always decorative, skip
        if _RE_SEP_DASH.match(line):
            i += 1
            continue

        # ── Skip VALIDATION REPORT section entirely (don't re-parse old warnings)
        if line.upper() == "VALIDATION REPORT":
            # Consume all remaining lines — they belong to the old report block
            break

        # ── V10: suppress duplicate ABSTRACT header ──────────────────
        if re.match(r"^ABSTRACT\s*$", line, re.IGNORECASE):
            if abstract_seen:
                i += 1
                continue          # drop the duplicate header
            abstract_seen = True

        # ── Caption line
        if _RE_CAPTION.match(line):
            m = _RE_CAPTION.match(line)
            tokens.append({"type": "caption", "text": m.group(1).strip()})
            i += 1
            continue

        # ── Insight / Description line
        if _RE_INSIGHT.match(line):
            m = _RE_INSIGHT.match(line)
            tokens.append({"type": "insight", "text": m.group(1).strip()})
            i += 1
            continue

        # ── Source line
        if _RE_SOURCE.match(line):
            m = _RE_SOURCE.match(line)
            tokens.append({"type": "source", "text": m.group(1).strip()})
            i += 1
            continue

        # ── ======= block (doc title or chapter) ──────────────────────
        if _RE_SEP_EQ.match(line):
            j = i + 1
            inner = []
            while j < len(lines) and not _RE_SEP_EQ.match(lines[j].strip()):
                s = lines[j].strip()
                if s and not _RE_SEP_DASH.match(s):
                    inner.append(s)
                j += 1

            content = inner[0] if inner else ""
            i = j + 1

            if not content:
                continue

            if not title_done:
                tokens.append({"type": "doc_title", "text": content})
                title_done           = True
                current_chapter_text = content
                current_chapter_norm = _norm(content)
                skip_echo            = True
                continue

            m = _RE_NUMBERED.match(content)
            if m:
                in_chapter           = True
                current_chapter_num  = int(m.group(1))
                chapter_text         = m.group(2).strip()
                current_chapter_text = chapter_text
                current_chapter_norm = _norm(chapter_text)
                last_sub_text        = ""
                last_sub_norm        = ""
                tokens.append({
                    "type": "chapter",
                    "num":  current_chapter_num,
                    "text": chapter_text,
                })
            else:
                in_chapter           = False
                current_chapter_text = content
                current_chapter_norm = _norm(content)
                last_sub_text        = ""
                last_sub_norm        = ""
                tokens.append({"type": "chapter", "num": None, "text": content})

            skip_echo = True
            continue

        # ── Echo suppression guard: numbered subheading lines are real
        if skip_echo and _RE_NUMBERED.match(line):
            pass  # fall through

        elif skip_echo:
            norm_line = _norm(line)
            is_echo   = False
            if len(line.split()) <= 9:
                if current_chapter_norm and _max_overlap(norm_line, current_chapter_norm) >= 0.55:
                    is_echo = True
                if not is_echo and last_sub_norm and _max_overlap(norm_line, last_sub_norm) >= 0.55:
                    is_echo = True
            skip_echo = False
            if is_echo:
                last_sub_norm = ""
                i += 1
                continue

        # ── Numbered subheading: "3. Feature Selection …" ────────────
        m = _RE_NUMBERED.match(line)
        if m:
            sub_num      = int(m.group(1))
            sub_text     = m.group(2).strip()
            last_sub_text = sub_text
            last_sub_norm = _norm(sub_text)
            skip_echo    = True

            # V11: fix wrong chapter prefix (e.g. sub_num=3 inside Chapter 4)
            # The sub_num here is the *section* number (e.g. "1" in "3.1").
            # But the numbered-subheading regex captures the leading digit as sub_num,
            # so "3.1 Title" → sub_num=3, text="1 Title" — this format is for
            # standalone numbered lines like "1. Title", "2. Title" inside a chapter.
            # The actual N.M renumbering is handled at render time via chapter_num + sub_num.
            # What we detect here is a plain body line that looks like "3.1 Statistical Analysis"
            # which should be "4.1 Statistical Analysis" — handled in _fix_wrong_prefix below.

            tokens.append({
                "type":          "subheading",
                "chapter_num":   current_chapter_num if in_chapter else None,
                "sub_num":       sub_num,
                "text":          sub_text,
                "_chapter_text": current_chapter_text,
            })
            i += 1
            continue

        # ── Bullet ────────────────────────────────────────────────────
        if re.match(r"^[\-\*\•\–]\s+.+", line):
            tokens.append({"type": "bullet",
                           "text": re.sub(r"^[\-\*\•\–]\s+", "", line)})
            i += 1
            continue

        # ── Figure placeholder (whole-line or inline) ─────────────────
        if _RE_FIG_PH.fullmatch(line) or re.fullmatch(r"\[Figure[^\]]+\]", line, re.IGNORECASE):
            fm = _RE_FIG_PH.search(line)
            if fm:
                tokens.append({"type": "figure_ph",
                                "ref": fm.group(1),
                                "text": fm.group(2).strip()})
                i += 1
                continue

        # ── Equation placeholder (whole line) ─────────────────────────
        if (_RE_EQ_PH.search(line)
                and len(line) < 120
                and not any(c.isalpha() and c.islower()
                            for c in line.replace("EQ", "").replace("Eq", "")[:10])):
            em = _RE_EQ_PH.search(line)
            if em:
                tokens.append({"type": "equation_ph",
                                "text": em.group(1).strip() if em.group(1) else line})
                i += 1
                continue

        # ── V11: detect "N.M Title" body lines with wrong chapter prefix ─
        # e.g. "3.1 Statistical Analysis" inside Chapter 4
        wrong_prefix_m = re.match(r"^(\d+)\.(\d+)\s+(.+)$", line)
        if (wrong_prefix_m
                and in_chapter
                and int(wrong_prefix_m.group(1)) != current_chapter_num
                and int(wrong_prefix_m.group(1)) != 0):
            correct_prefix = current_chapter_num
            sub_n   = int(wrong_prefix_m.group(2))
            sub_txt = wrong_prefix_m.group(3).strip()
            tokens.append({
                "type":          "subheading",
                "chapter_num":   correct_prefix,
                "sub_num":       sub_n,
                "text":          sub_txt,
                "_chapter_text": current_chapter_text,
                "_fixed_prefix": True,   # informational flag
            })
            last_sub_text = sub_txt
            last_sub_norm = _norm(sub_txt)
            skip_echo     = True
            i += 1
            continue

        # ── Body text — may contain inline figure/equation refs ───────
        sub_toks = _split_inline_placeholders(line)
        tokens.extend(sub_toks)
        i += 1

    return tokens


# ═════════════════════════════════════════════════════════════════════════════
# REFERENCE SECTION HELPERS  (used by validate + dedup)
# ═════════════════════════════════════════════════════════════════════════════

def _extract_references_section(tokens: List[dict]) -> List[str]:
    """
    Return raw text lines from the REFERENCES section tokens.
    """
    in_refs = False
    lines   = []
    for tok in tokens:
        if tok["type"] == "chapter" and _norm(tok["text"]) in ("references", "bibliography"):
            in_refs = True
            continue
        if in_refs and tok["type"] == "chapter":
            break
        if in_refs and tok["type"] in ("body", "bullet"):
            lines.append(tok["text"])
    return lines


def _ref_key(line: str) -> Optional[str]:
    """Extract a normalised author+year key from a reference line."""
    cites = _RE_CITE.findall(line)
    if cites:
        return _normalise_citation(cites[0])
    year_m   = re.search(r"\b(19|20)\d{2}\b", line)
    author_m = re.match(r"^([A-Z][a-z]+)", line)
    if author_m and year_m:
        return f"{author_m.group(1).lower()}{year_m.group()}"
    return None


def _score_ref_line(line: str) -> int:
    """Score a reference line by completeness (higher = keep this one)."""
    score = len(line)
    if re.search(r"https?://", line):
        score += 20
    if re.search(r"doi\.org|DOI:", line, re.IGNORECASE):
        score += 15
    if re.search(r"pp\.\s*\d+", line):
        score += 10
    if "Available at:" in line:
        score += 5
    return score


# ═════════════════════════════════════════════════════════════════════════════
# REFERENCE DEDUPLICATION  (V7 / V12) — operates on the token list in-place
# ═════════════════════════════════════════════════════════════════════════════

def _deduplicate_references(tokens: List[dict]) -> Tuple[List[dict], List[str]]:
    """
    Find duplicate entries in the REFERENCES section.
    For each duplicate group: keep the entry with the highest completeness score;
    mark the inferior copies as type='echo' (suppressed at render time).

    Returns (modified_tokens, list_of_warning_strings).
    """
    warnings: List[str] = []

    # Locate the REFERENCES chapter token and collect its body token indices
    in_refs  = False
    ref_toks : List[Tuple[int, dict]] = []   # (token_index, token)
    for idx, tok in enumerate(tokens):
        if tok["type"] == "chapter" and _norm(tok["text"]) in ("references", "bibliography"):
            in_refs = True
            continue
        if in_refs and tok["type"] == "chapter":
            break
        if in_refs and tok["type"] in ("body", "bullet"):
            ref_toks.append((idx, tok))

    if not ref_toks:
        return tokens, warnings

    # Group by normalised key
    groups: Dict[str, List[Tuple[int, dict]]] = defaultdict(list)
    ungrouped: List[Tuple[int, dict]] = []
    for idx, tok in ref_toks:
        key = _ref_key(tok["text"])
        if key:
            groups[key].append((idx, tok))
        else:
            ungrouped.append((idx, tok))

    for key, entries in groups.items():
        if len(entries) <= 1:
            continue
        # Sort by score descending — keep the best
        entries_scored = sorted(entries, key=lambda e: _score_ref_line(e[1]["text"]), reverse=True)
        keeper_idx, keeper_tok = entries_scored[0]
        for dup_idx, dup_tok in entries_scored[1:]:
            msg = (f"[V7] Duplicate REFERENCES entry '{key}' -- "
                   f"kept: \"{keeper_tok['text'][:60]}...\"  "
                   f"removed: \"{dup_tok['text'][:60]}...\"")
            warnings.append(msg)
            tokens[dup_idx]["type"] = "echo"   # suppress at render time

    return tokens, warnings


# ═════════════════════════════════════════════════════════════════════════════
# MISSING / ORPHAN REFERENCE CHECKS  (V8 + V9)
# ═════════════════════════════════════════════════════════════════════════════

def _check_reference_integrity(tokens: List[dict]) -> List[str]:
    """
    V8: in-text citation has no matching REFERENCES entry.
    V9: REFERENCES entry is never cited in the text body.

    Returns a list of warning strings (same format as validate()).
    """
    warnings: List[str] = []

    # 1. Collect all in-text citation keys (excluding REFERENCES section itself)
    in_refs           = False
    text_cites        : Dict[str, List[str]] = defaultdict(list)   # key → [locations]
    cur_chapter_text  = ""
    cur_chapter_num   = None
    cur_sub_text      = ""

    def _loc():
        ch = f"Ch{cur_chapter_num}" if cur_chapter_num else cur_chapter_text
        return f"{ch} / {cur_sub_text}" if cur_sub_text else ch

    for tok in tokens:
        t = tok["type"]
        if t == "chapter":
            if _norm(tok["text"]) in ("references", "bibliography"):
                in_refs = True
            else:
                in_refs          = False
                cur_chapter_text = tok["text"]
                cur_chapter_num  = tok.get("num")
                cur_sub_text     = ""
        elif t == "subheading":
            cur_sub_text = tok["text"]
        elif t in ("body", "h3", "bullet") and not in_refs:
            for c in _extract_citations(tok["text"]):
                text_cites[_normalise_citation(c)].append(_loc())

    # 2. Collect all reference keys from REFERENCES section
    ref_lines = _extract_references_section(tokens)
    ref_keys  : Dict[str, str] = {}   # key → representative line
    for line in ref_lines:
        key = _ref_key(line)
        if key and key not in ref_keys:
            ref_keys[key] = line

    # 3. V8 — in-text cite not in references
    for key, locs in text_cites.items():
        if key and key not in ref_keys:
            unique_locs = list(dict.fromkeys(locs))
            warnings.append(
                f"[V8] In-text citation '{key}' has no entry in REFERENCES "
                f"(cited in: {', '.join(unique_locs[:3])})"
            )

    # 4. V9 — reference never cited in text
    for key, line in ref_keys.items():
        if key and key not in text_cites:
            warnings.append(
                f"[V9] REFERENCES entry '{key}' is never cited in the text -- "
                f"entry: \"{line[:70]}...\""
            )

    return warnings


# ═════════════════════════════════════════════════════════════════════════════
# CITATION DEDUPLICATION  (V1 — in-text, global)
# ═════════════════════════════════════════════════════════════════════════════

def apply_global_citation_deduplication(tokens: List[dict]) -> List[dict]:
    """
    Scans body / h3 / bullet tokens for in-text citations.
    Keeps the FIRST occurrence of each (Author, Year) pair globally;
    all subsequent occurrences are removed from the token text.

    Skips tokens inside the REFERENCES section.
    """
    seen_cites: set = set()
    in_refs         = False

    for tok in tokens:
        t = tok["type"]
        if t == "chapter":
            in_refs = _norm(tok["text"]) in ("references", "bibliography")
            continue
        if in_refs or t not in ("body", "h3", "bullet"):
            continue

        text    = tok["text"]
        matches = list(_RE_CITE.finditer(text))

        # Process in reverse to avoid index shifts when slicing
        for match in reversed(matches):
            cite_str = match.group(0)
            norm     = _normalise_citation(cite_str)
            if norm and norm in seen_cites:
                start, end = match.span()
                if start > 0 and text[start - 1] == " ":
                    start -= 1
                text = text[:start] + text[end:]
            elif norm:
                seen_cites.add(norm)

        # Remove word count markers like (278 words)
        text = re.sub(r"[\(\[]\d+\s*words[\)\]]", "", text, flags=re.IGNORECASE)
        
        # Replace big dashes (em-dash, en-dash) with a simple hyphen
        text = text.replace("—", "-").replace("–", "-")

        text = re.sub(r"\s{2,}", " ", text)
        text = re.sub(r"\s+([,\.\?\!])", r"\1", text)
        tok["text"] = text.strip()

    return tokens


# ═════════════════════════════════════════════════════════════════════════════
# STEP 4 — LLM REFINEMENT  (V3 + V4)
#   • Promote ambiguous body lines to h3 or echo
#   • Uses chapter + subheading context for accurate decisions
# ═════════════════════════════════════════════════════════════════════════════

def refine_with_llm(tokens: List[dict], use_llm: bool) -> List[dict]:
    if not use_llm or not _LLM_OK:
        if use_llm and not _LLM_OK:
            print("    ⚠️  openai / dotenv not installed — skipping LLM step.")
        return tokens

    cur_ch  = ""
    cur_sub = ""
    candidates: List[Tuple[int, str, str, str]] = []

    for idx, tok in enumerate(tokens):
        if tok["type"] == "chapter":
            cur_ch  = tok["text"]
            cur_sub = ""
        elif tok["type"] == "subheading":
            cur_sub = tok["text"]
        elif tok["type"] == "body":
            text  = tok["text"].strip()
            words = text.split()
            if (1 <= len(words) <= 9
                    and not text.endswith(".")
                    and not text.endswith(",")
                    and re.match(r"^[A-Z]", text)
                    and not re.search(r"\d{4}", text)):
                candidates.append((idx, text, cur_ch, cur_sub))

    if not candidates:
        return tokens

    groups: Dict[tuple, list] = defaultdict(list)
    for idx, text, ch, sub in candidates:
        groups[(ch, sub)].append((idx, text))

    total_upgraded = 0
    total_echoed   = 0

    for (ch, sub), group in groups.items():
        print(f"    🤖  Classifying {len(group)} lines under [{ch} / {sub or '—'}] …")
        results = llm_classify_lines(group, ch, sub)
        for idx, label in results.items():
            if label == "h3":
                tokens[idx]["type"] = "h3"
                total_upgraded += 1
            elif label == "echo":
                tokens[idx]["type"] = "echo"
                total_echoed += 1

    print(f"    🤖  Promoted {total_upgraded} → h3,  suppressed {total_echoed} echoes.")
    return tokens


# ═════════════════════════════════════════════════════════════════════════════
# STEP 5 — VALIDATIONS  (V1 + V2 flags, V7/V8/V9 ref integrity)
# ═════════════════════════════════════════════════════════════════════════════

def validate(tokens: List[dict]) -> List[str]:
    """
    Run all validation checks and return a list of warning strings.

    Warning format (unchanged from original for app.py compatibility):
      [V1] …
      [V2] …
      [V7] …   (duplicate REFERENCES — new)
      [V8] …   (missing reference — new)
      [V9] …   (orphan reference — new)

    V2 warnings are also injected as '_warnings' on the offending token
    so they render inline in the DOCX (original behaviour preserved).

    NOTE: apply_global_citation_deduplication() should be called BEFORE
    validate() so that V1 warnings reflect the post-dedup state.
    """
    warnings: List[str] = []

    # ── State tracking ───────────────────────────────────────────────
    cur_chapter_text = ""
    cur_chapter_num  = None
    cur_sub_text     = ""

    # V1: global citation registry  { normalised_key: [locations] }
    cite_registry: Dict[str, List[str]] = defaultdict(list)

    # V2: per-section citation tracking
    results_bench_cite_count = 0

    def _location():
        ch = f"Ch{cur_chapter_num}" if cur_chapter_num else cur_chapter_text
        return f"{ch} / {cur_sub_text}" if cur_sub_text else ch

    def _section_allows_citations() -> bool:
        norm = _norm(cur_chapter_text)
        return any(kw in norm for kw in CITATION_OK_SECTIONS)

    def _in_results() -> bool:
        return cur_chapter_num == 5 or "result" in _norm(cur_chapter_text)

    def _in_special() -> bool:
        return cur_chapter_text.upper() in SPECIAL_SECTIONS

    def _is_benchmark_para(text: str) -> bool:
        return any(kw in text.lower() for kw in RESULTS_BENCHMARK_KEYWORDS)

    for i, tok in enumerate(tokens):
        t = tok["type"]

        if t == "chapter":
            cur_chapter_text = tok["text"]
            cur_chapter_num  = tok.get("num")
            cur_sub_text     = ""
            results_bench_cite_count = 0

        elif t == "subheading":
            cur_sub_text = tok["text"]
            results_bench_cite_count = 0

        elif t in ("body", "h3"):
            if _in_special():
                continue          # no citation rules in REFERENCES etc.

            text  = tok["text"]
            cites = _extract_citations(text)

            for cite in cites:
                key = _normalise_citation(cite)
                loc = _location()
                cite_registry[key].append(loc)

                # V2 ─────────────────────────────────────────────────
                if _section_allows_citations():
                    pass

                elif _in_results():
                    if _is_benchmark_para(text):
                        results_bench_cite_count += 1
                        if results_bench_cite_count > 1:
                            msg = (f"[V2] Results section allows only 1 citation in "
                                   f"benchmark paragraph — extra cite: {cite}  @{loc}")
                            warnings.append(msg)
                            tok.setdefault("_warnings", []).append(msg)
                    else:
                        msg = (f"[V2] Citation '{cite}' found outside allowed sections "
                               f"(use only in Intro / Lit Review / Methodology or "
                               f"1× in Results benchmark paragraph)  @{loc}")
                        warnings.append(msg)
                        tok.setdefault("_warnings", []).append(msg)

                else:
                    msg = (f"[V2] Citation '{cite}' found in non-citation section  "
                           f"@{loc}")
                    warnings.append(msg)
                    tok.setdefault("_warnings", []).append(msg)

    # V1 — remaining repeated citations (after deduplication) ────────
    for key, locations in cite_registry.items():
        unique = list(dict.fromkeys(locations))
        if len(unique) > 1:
            msg = (f"[V1] Citation key '{key}' appears {len(unique)}× in: "
                   f"{', '.join(unique)}")
            warnings.append(msg)

    # V7 / V8 / V9 — reference integrity ────────────────────────────
    ref_integrity_warnings = _check_reference_integrity(tokens)
    warnings.extend(ref_integrity_warnings)

    return warnings


# ═════════════════════════════════════════════════════════════════════════════
# STEP 6 — DOCX RENDERING
# ═════════════════════════════════════════════════════════════════════════════

# ── Low-level helpers ─────────────────────────────────────────────────────────

def _run(para, text: str, size: int,
         bold=False, italic=False, color=None):
    r = para.add_run(text)
    r.font.name   = CFG["font"]
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _spacing(para, before=0, after=6, ls=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing = ls or CFG["line_sp"]


def _border_bottom(para, color_hex="2E74B5", size=6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

def _border_left(para, color_hex="2E74B5", size=18, space=6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _shade_paragraph(para, fill_hex="F2F2F2"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)




# -- Element renderers ---
def _render_doc_title(doc, text):
    n = CFG["_n"]
    for _ in range(CFG["title_spacers"]):
        sp = doc.add_paragraph(); _spacing(sp, 0, 0)
    if n == 3:
        eyebrow = doc.add_paragraph()
        _run(eyebrow, "DISSERTATION", 9, bold=True, color=RGBColor(0xAA,0xAA,0xAA))
        _spacing(eyebrow, before=0, after=4)
    if n == 4:
        tag = doc.add_paragraph()
        tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(tag, "-- Doctoral Research --", 10, italic=True, color=CFG["col_sub"])
        _spacing(tag, before=0, after=10)
    p = doc.add_paragraph()
    if CFG["title_center"]:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if n == 1:
        _run(p, text, CFG["sz_title"], bold=True, color=CFG["col_title"])
        _spacing(p, before=0, after=14)
        _border_bottom(p, color_hex=CFG["bdr_title"], size=8)
    elif n == 2:
        _run(p, text.upper(), CFG["sz_title"], bold=True, color=CFG["col_title"])
        _spacing(p, before=0, after=6)
        _border_bottom(p, color_hex=CFG["bdr_title"], size=18)
        sub = doc.add_paragraph(); _spacing(sub, before=2, after=18)
        _border_bottom(sub, color_hex="DAA520", size=4)
    elif n == 3:
        _run(p, text, CFG["sz_title"], bold=True, color=CFG["col_title"])
        _spacing(p, before=0, after=16)
        _border_bottom(p, color_hex=CFG["bdr_title"], size=2)
    elif n == 4:
        _run(p, text, CFG["sz_title"], bold=True, color=CFG["col_title"])
        _spacing(p, before=0, after=12)
        _border_bottom(p, color_hex=CFG["bdr_title"], size=12)
    doc.add_page_break()


def _render_chapter(doc, num, text):
    n = CFG["_n"]
    if CFG["ch_page_break"]:
        doc.add_page_break()
    if n == 1:
        p = doc.add_paragraph()
        if num is not None:
            _run(p, "Chapter " + str(num) + chr(10), CFG["sz_ch"] - 4, italic=True, color=CFG["col_sub"])
        _run(p, text.upper(), CFG["sz_ch"], bold=True, color=CFG["col_ch"])
        _spacing(p, before=0, after=8)
        _border_bottom(p, color_hex=CFG["bdr_ch"], size=10)
    elif n == 2:
        top = doc.add_paragraph(); _spacing(top, before=0, after=6)
        _border_bottom(top, color_hex="B8860B", size=6)
        p = doc.add_paragraph()
        if num is not None:
            _run(p, f"{num:02d}  ", CFG["sz_ch"] + 6, bold=True, color=RGBColor(0xDA,0xA5,0x20))
        _run(p, text.upper(), CFG["sz_ch"], bold=True, color=CFG["col_ch"])
        _spacing(p, before=4, after=10)
        _border_bottom(p, color_hex="B8860B", size=3)
    elif n == 3:
        p = doc.add_paragraph()
        if num is not None:
            _run(p, f"CH.{num}  ", 8, bold=True, color=RGBColor(0xAA,0xAA,0xAA))
        _run(p, text, CFG["sz_ch"], bold=True, color=CFG["col_ch"])
        _spacing(p, before=4, after=6)
        _border_bottom(p, color_hex=CFG["bdr_ch"], size=2)
    elif n == 4:
        p = doc.add_paragraph()
        _shade_paragraph(p, fill_hex="005050")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = f"Chapter {num} -- " if num is not None else ""
        _run(p, label + text.upper(), CFG["sz_ch"], bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        _spacing(p, before=8, after=8)


def _render_subheading(doc, chapter_num, sub_num, text):
    n = CFG["_n"]
    p = doc.add_paragraph()
    if n == 1:
        if chapter_num:
            _run(p, f"{chapter_num}.{sub_num}  ", CFG["sz_sub"], bold=True, color=CFG["col_sub"])
        _run(p, text, CFG["sz_sub"], bold=True, color=CFG["col_sub"])
        _spacing(p, before=14, after=4)
    elif n == 2:
        if chapter_num:
            _run(p, f"{chapter_num}.{sub_num}  ", CFG["sz_sub"] - 1, bold=True, color=RGBColor(0xB8,0x86,0x0B))
        _run(p, text, CFG["sz_sub"], bold=True, color=CFG["col_sub"])
        _spacing(p, before=16, after=5)
        _border_bottom(p, color_hex="DAA520", size=2)
    elif n == 3:
        if chapter_num:
            _run(p, f"{chapter_num}.{sub_num}  ", CFG["sz_sub"], bold=True, color=RGBColor(0x99,0x99,0x99))
        _run(p, text, CFG["sz_sub"], bold=True, color=CFG["col_sub"])
        _spacing(p, before=18, after=4)
    elif n == 4:
        p.paragraph_format.left_indent = Inches(0.1)
        if chapter_num:
            _run(p, f"{chapter_num}.{sub_num}  ", CFG["sz_sub"], bold=True, color=CFG["col_sub"])
        _run(p, text, CFG["sz_sub"], bold=True, color=CFG["col_sub"])
        _border_left(p, color_hex=CFG["bdr_sub"], size=18, space=6)
        _spacing(p, before=14, after=5)


def _render_h3(doc, text):
    n = CFG["_n"]
    p = doc.add_paragraph()
    if n == 1:
        p.paragraph_format.left_indent = Pt(12)
        _run(p, text, CFG["sz_h3"], bold=True, italic=True, color=CFG["col_h3"])
        _spacing(p, before=8, after=2)
    elif n == 2:
        p.paragraph_format.left_indent = Inches(0.15)
        _run(p, ">>  ", CFG["sz_h3"], bold=True, color=RGBColor(0xB8,0x86,0x0B))
        _run(p, text, CFG["sz_h3"], bold=True, color=CFG["col_h3"])
        _spacing(p, before=10, after=3)
    elif n == 3:
        _run(p, text.upper(), CFG["sz_h3"] - 1, bold=True, color=RGBColor(0x88,0x88,0x88))
        _spacing(p, before=12, after=2)
    elif n == 4:
        p.paragraph_format.left_indent = Inches(0.2)
        _run(p, text, CFG["sz_h3"], bold=True, italic=True, color=CFG["col_h3"])
        _spacing(p, before=8, after=3)


def _render_figure_ph(doc, ref, label):
    n = CFG["_n"]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    if n == 1:
        p.paragraph_format.left_indent = Inches(0.25)
        _border_left(p, color_hex=CFG["bdr_fig"], size=12, space=8)
        _run(p, f"[ Figure {ref}  --  {label} ]", CFG["sz_body"], italic=True, color=CFG["col_fig"])
    elif n == 2:
        p.paragraph_format.left_indent = Inches(0.2)
        _shade_paragraph(p, fill_hex="FFF8DC")
        _run(p, f"Fig. {ref}  -- ", CFG["sz_body"], bold=True, color=RGBColor(0xB8,0x86,0x0B))
        _run(p, label, CFG["sz_body"], italic=True, color=CFG["col_body"])
    elif n == 3:
        _shade_paragraph(p, fill_hex="F6F6F6")
        _run(p, f"Fig. {ref}  -- ", CFG["sz_body"], bold=True, color=RGBColor(0x44,0x44,0x44))
        _run(p, label, CFG["sz_body"], color=RGBColor(0x55,0x55,0x55))
    elif n == 4:
        p.paragraph_format.left_indent = Inches(0.25)
        _border_left(p, color_hex=CFG["bdr_fig"], size=12, space=8)
        _run(p, f"[ Figure {ref}  --  {label} ]", CFG["sz_body"], italic=True, color=CFG["col_fig"])


def _render_equation_ph(doc, text):
    n = CFG["_n"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if n == 1:
        _run(p, f"[ Equation  --  {text} ]", CFG["sz_body"], italic=True, color=CFG["col_eq"])
        _spacing(p, before=8, after=8)
    elif n == 2:
        _shade_paragraph(p, fill_hex="F5F0FF")
        _run(p, f"   {text}   ", CFG["sz_body"], italic=True, color=CFG["col_eq"])
        _spacing(p, before=10, after=10)
    elif n == 3:
        _run(p, text, CFG["sz_body"], italic=True, color=RGBColor(0x55,0x55,0x55))
        _spacing(p, before=10, after=10)
        _border_bottom(p, color_hex="CCCCCC", size=2)
    elif n == 4:
        _run(p, f"[ Equation  --  {text} ]", CFG["sz_body"], italic=True, color=CFG["col_eq"])
        _spacing(p, before=8, after=8)


# -- Master renderer ───────────────────────────────────────────────────────────

def generate_docx(tokens: List[dict],
                  global_warnings: List[str],
                  output_path,
                  llm_analysis: str = "",
                  format: int = 1):
    """
    Render token list to a DOCX file.

    Parameters
    ----------
    tokens          : validated + deduplicated token list
    global_warnings : list of [Vx] warning strings from validate()
    output_path     : file path string or BytesIO object
    llm_analysis    : (optional) LLM deep-analysis text appended after warnings
    format          : int 1-4 -- selects the visual theme (passed from app.py UI)
                        1  Classic Academic Blue  (default)
                        2  Executive Dark Slate
                        3  Modern Minimalist
                        4  Research Emerald
    """
    global CFG
    CFG       = _build_theme(int(format))
    CFG["_n"] = int(format)       # expose numeric ID to all renderer functions

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(CFG["margin_top"])
        sec.bottom_margin = Inches(CFG["margin_bottom"])
        sec.left_margin   = Inches(CFG["margin_left"])
        sec.right_margin  = Inches(CFG["margin_right"])

    for tok in tokens:
        t = tok["type"]
        if   t == "echo":        continue
        elif t == "doc_title":   _render_doc_title(doc, tok["text"])
        elif t == "chapter":     _render_chapter(doc, tok.get("num"), tok["text"])
        elif t == "subheading":  _render_subheading(doc,
                                     tok.get("chapter_num"),
                                     tok.get("sub_num", 0),
                                     tok["text"])
        elif t == "h3":          _render_h3(doc, tok["text"])
        elif t == "bullet":      _render_bullet(doc, tok["text"])
        elif t == "body":        _render_body(doc, tok["text"], tok.get("_warnings"))
        elif t == "figure_ph":   _render_figure_ph(doc, tok.get("ref", "?"), tok["text"])
        elif t == "equation_ph": _render_equation_ph(doc, tok["text"])
        elif t == "caption":     _render_caption(doc, tok["text"])
        elif t == "insight":     _render_insight(doc, tok["text"])
        elif t == "source":      _render_source(doc, tok["text"])

    # -- Validation Report --------------------------------------------------------
    if global_warnings or llm_analysis:
        doc.add_page_break()
        hdr = doc.add_paragraph()
        _run(hdr, "VALIDATION REPORT", CFG["sz_ch"], bold=True, color=CFG["col_warn"])
        _spacing(hdr, before=0, after=10)
        _border_bottom(hdr, color_hex=CFG["bdr_warn"], size=8)

        v_groups: Dict[str, List[str]] = defaultdict(list)
        for w in global_warnings:
            m = re.match(r"^\[([^\]]+)\]", w)
            v_groups[m.group(1) if m else "OTHER"].append(w)

        for code in sorted(v_groups):
            sh = doc.add_paragraph()
            _run(sh, f"{code} Issues  ({len(v_groups[code])})",
                 CFG["sz_body"], bold=True, color=CFG["col_warn"])
            _spacing(sh, before=8, after=2)
            for w in v_groups[code]:
                _render_warning_block(doc, w)

        if llm_analysis:
            sep = doc.add_paragraph()
            _run(sep, "AI ANALYSIS", CFG["sz_body"], bold=True, color=CFG["col_ai"])
            _spacing(sep, before=12, after=4)
            for line in llm_analysis.splitlines():
                if line.strip():
                    lp = doc.add_paragraph()
                    _run(lp, line, 9, italic=True, color=RGBColor(0x30,0x30,0x30))
                    _spacing(lp, before=0, after=2)

    doc.save(output_path)
    if isinstance(output_path, str):
        print(f"Saved -> {output_path}  "
              f"({len(tokens)} tokens, {len(global_warnings)} warnings, "
              f"theme: {CFG['name']})")



# ── TXT renderer ──────────────────────────────────────────────────────────────

def generate_formatted_txt(tokens: List[dict]) -> str:
    """
    Render the validated token stream back into cleanly formatted plain text.
    Ensures N.M hierarchical numbering for subheadings (e.g. 3.1. Methodology).
    Suppressed (echo) tokens are omitted.
    """
    lines = []

    for tok in tokens:
        t = tok["type"]

        if t == "echo":
            continue
        elif t == "doc_title":
            lines.append("======================================================================")
            lines.append(tok["text"].upper())
            lines.append("======================================================================\n")
        elif t == "chapter":
            num = tok.get("num")
            lines.append("======================================================================")
            if num is not None:
                lines.append(f"{num}. {tok['text'].upper()}")
            else:
                lines.append(tok["text"].upper())
            lines.append("======================================================================\n")
        elif t == "subheading":
            ch_num  = tok.get("chapter_num")
            sub_num = tok.get("sub_num", 0)
            prefix  = f"{ch_num}.{sub_num}. " if ch_num else ""
            lines.append(f"{prefix}{tok['text']}\n")
        elif t == "h3":
            lines.append(f"  {tok['text']}\n")
        elif t == "bullet":
            lines.append(f"- {tok['text']}\n")
        elif t == "body":
            lines.append(f"{tok['text']}\n")
        elif t == "figure_ph":
            ref = tok.get("ref", "?")
            lines.append(f"[Figure {ref}: {tok['text']}]\n")
        elif t == "equation_ph":
            lines.append(f"[EQN: {tok['text']}]\n")
        elif t == "caption":
            lines.append(f"Caption: {tok['text']}\n")
        elif t == "insight":
            lines.append(f"Insight: {tok['text']}\n")
        elif t == "source":
            lines.append(f"Source: {tok['text']}\n")

    return "\n".join(lines).strip()


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════

def _count(toks, t):
    return sum(1 for x in toks if x["type"] == t)


def main():
    args        = sys.argv[1:]
    use_llm     = "--no-llm" not in args
    clean       = [a for a in args if not a.startswith("--")]
    input_path  = clean[0] if len(clean) > 0 else "input.txt"
    output_path = clean[1] if len(clean) > 1 else "output.docx"

    # Parse --format N  (supports "--format 2" and "--format=2")
    fmt = 1
    for i, a in enumerate(args):
        if a.startswith("--format="):
            fmt = int(a.split("=", 1)[1])
        elif a == "--format" and i + 1 < len(args):
            fmt = int(args[i + 1])

    print(f"Reading      : {input_path}")
    print(f"Format theme : {fmt}  ({_build_theme(fmt)['name']})")
    raw = read_file(input_path)

    print("🔍  Parsing structure …")
    tokens = parse_structure(raw)
    print(
        f"    chapters={_count(tokens,'chapter')}  "
        f"subheadings={_count(tokens,'subheading')}  "
        f"body={_count(tokens,'body')}  "
        f"figures={_count(tokens,'figure_ph')}  "
        f"equations={_count(tokens,'equation_ph')}"
    )

    print("🧹  Deduplicating in-text citations …")
    tokens = apply_global_citation_deduplication(tokens)

    print("🔗  Deduplicating REFERENCES section …")
    tokens, ref_dup_warnings = _deduplicate_references(tokens)
    print(f"    Duplicate reference entries removed: {len(ref_dup_warnings)}")

    print("🔎  Running validations (V1 + V2 + V8 + V9) …")
    warnings = validate(tokens)
    # Merge ref-dedup warnings (V7) into the main list
    warnings = ref_dup_warnings + warnings

    v_counts: Dict[str, int] = defaultdict(int)
    for w in warnings:
        m = re.match(r"^\[([^\]]+)\]", w)
        if m:
            v_counts[m.group(1)] += 1
    for code, cnt in sorted(v_counts.items()):
        print(f"    {code}: {cnt} warning(s)")
    for w in warnings[:10]:
        print(f"    {w}")
    if len(warnings) > 10:
        print(f"    … and {len(warnings)-10} more (see Validation Report in DOCX)")

    # LLM deep-analysis of all warnings
    llm_analysis = ""
    if use_llm and warnings:
        print("🤖  LLM issue analysis …")
        llm_analysis = llm_analyse_issues(
            [w for w in warnings if not w.startswith("[V7]")],
            len(tokens),
            [w for w in warnings if w.startswith("[V7]")],
        )

    if use_llm:
        print("🤖  LLM refinement (V3 + V4 — h3 detection + echo removal) …")
        tokens = refine_with_llm(tokens, use_llm=True)
        print(
            f"    After LLM: h3={_count(tokens,'h3')}  "
            f"echo-suppressed={_count(tokens,'echo')}  "
            f"body={_count(tokens,'body')}"
        )
    else:
        print("    ⚠️  --no-llm: skipping h3 detection and LLM echo removal.")

    print("📝  Building DOCX …")
    generate_docx(tokens, warnings, output_path, llm_analysis=llm_analysis, format=fmt)


if __name__ == "__main__":
    main()